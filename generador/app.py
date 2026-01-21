"""
APLICACIÓN WEB: GENERADOR DE DOCUMENTOS CIENTÍFICOS
Autor: Ismael Antonio Cárdenas López
Licenciado en Matemática - UNAN León
Universidad Nacional Padre Gaspar García Laviana
"""

import os
import re
import tempfile
import uuid
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, send_file, jsonify, send_from_directory
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# CONFIGURACIÓN DE LA APLICACIÓN
# ============================================================================

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)
app.secret_key = os.environ.get('SECRET_KEY', 'clave-secreta-documentos-cientificos')

# Crear directorios necesarios
for dir_path in ['temp', 'static/images', 'uploads']:
    Path(dir_path).mkdir(exist_ok=True)

# ============================================================================
# RUTAS PRINCIPALES
# ============================================================================

@app.route('/')
def index():
    """Página principal - Sirve el HTML estático"""
    return send_from_directory('.', 'index.html')

@app.route('/api/generar-word', methods=['POST'])
def generar_word():
    """API para generar documento Word"""
    try:
        data = request.json
        titulo = data.get('titulo', 'Documento Científico')
        contenido = data.get('contenido', '')
        autor = data.get('autor', 'Ismael Antonio Cárdenas López')
        
        # Crear documento Word
        doc = Document()
        
        # Configurar márgenes
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        
        # Portada
        titulo_parrafo = doc.add_paragraph()
        titulo_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo_parrafo.add_run(titulo.upper())
        titulo_run.font.size = Pt(28)
        titulo_run.font.color.rgb = RGBColor(0, 51, 102)
        titulo_run.font.bold = True
        titulo_run.font.name = 'Times New Roman'
        
        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph().add_run().add_break()
        
        # Información del autor
        autor_parrafo = doc.add_paragraph()
        autor_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        autor_run = autor_parrafo.add_run(autor)
        autor_run.font.size = Pt(14)
        autor_run.font.color.rgb = RGBColor(64, 64, 64)
        autor_run.font.bold = True
        
        # Contenido
        doc.add_page_break()
        doc.add_heading('CONTENIDO', 1)
        
        # Procesar contenido
        lineas = contenido.split('\n')
        for linea in lineas:
            if linea.strip():
                if linea.startswith('# '):
                    doc.add_heading(linea[2:], 1)
                elif linea.startswith('## '):
                    doc.add_heading(linea[3:], 2)
                elif linea.startswith('### '):
                    doc.add_heading(linea[4:], 3)
                else:
                    doc.add_paragraph(linea)
        
        # Guardar archivo temporal
        filename = f"documento_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join('temp', filename)
        doc.save(filepath)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'message': 'Documento Word generado exitosamente'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/descargar/<filename>')
def descargar_word(filename):
    """Descargar documento generado"""
    try:
        filepath = os.path.join('temp', filename)
        if not os.path.exists(filepath):
            return jsonify({'error': 'Archivo no encontrado'}), 404
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=f"documento_cientifico_{filename}",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/generar-latex', methods=['POST'])
def generar_latex():
    """API para generar código LaTeX"""
    try:
        data = request.json
        titulo = data.get('titulo', 'Documento Científico')
        contenido = data.get('contenido', '')
        autor = data.get('autor', 'Ismael Antonio Cárdenas López')
        
        # Plantilla LaTeX completa
        latex_template = f"""% ============================================
% DOCUMENTO CIENTÍFICO EN LATEX
% Generado automáticamente
% Autor: {autor}
% Fecha: {datetime.now().strftime('%d/%m/%Y')}
% ============================================

\\documentclass[12pt, a4paper]{{article}}

% Paquetes esenciales
\\usepackage{{amsmath, amssymb, amsthm, amsfonts}}
\\usepackage{{mathtools, bm, physics}}
\\usepackage{{geometry}}
\\geometry{{a4paper, margin=2.5cm}}
\\usepackage[spanish]{{babel}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{graphicx}}
\\usepackage{{hyperref}}
\\usepackage{{fancyhdr}}
\\usepackage{{titlesec}}

% Configuración
\\title{{{titulo}}}
\\author{{{autor} \\\\ Licenciado en Matemática \\\\ UNAN León}}
\\date{{\\today}}

\\begin{{document}}

\\maketitle

% Contenido
{contenido}

\\end{{document}}
"""
        
        # Guardar archivo .tex
        filename = f"documento_{datetime.now().strftime('%Y%m%d_%H%M%S')}.tex"
        filepath = os.path.join('temp', filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(latex_template)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'content': latex_template,
            'message': 'Código LaTeX generado exitosamente'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/limpiar-temp')
def limpiar_temp():
    """Limpiar archivos temporales antiguos"""
    try:
        import time
        import glob
        
        now = time.time()
        temp_files = glob.glob('temp/*')
        deleted = 0
        
        for filepath in temp_files:
            if os.path.isfile(filepath):
                # Eliminar archivos con más de 1 hora
                if now - os.path.getmtime(filepath) > 3600:
                    os.remove(filepath)
                    deleted += 1
        
        return jsonify({
            'success': True,
            'deleted': deleted,
            'message': f'Se eliminaron {deleted} archivos temporales'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/<path:path>')
def serve_static(path):
    """Servir archivos estáticos"""
    return send_from_directory('.', path)

# ============================================================================
# EJECUCIÓN
# ============================================================================

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    
    print("=" * 60)
    print("GENERADOR DE DOCUMENTOS CIENTÍFICOS")
    print("Autor: Ismael Antonio Cárdenas López")
    print("Licenciado en Matemática - UNAN León")
    print("=" * 60)
    print(f"Servidor ejecutándose en: http://localhost:{port}")
    print("=" * 60)
    
    app.run(host='0.0.0.0', port=port, debug=debug)